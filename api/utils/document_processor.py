from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Optional
import logging

logger = logging.getLogger(__name__)

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Add hyperlink to paragraph with formatting"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def copy_run_formatting(source_run, target_run):
    """Copy formatting from source to target run"""
    try:
        if hasattr(source_run, 'bold') and source_run.bold is not None:
            target_run.bold = source_run.bold
        if hasattr(source_run, 'italic') and source_run.italic is not None:
            target_run.italic = source_run.italic
        if hasattr(source_run, 'underline') and source_run.underline is not None:
            target_run.underline = source_run.underline
        if hasattr(source_run, 'font'):
            if hasattr(source_run.font, 'color') and source_run.font.color and source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            if hasattr(source_run.font, 'name') and source_run.font.name:
                target_run.font.name = source_run.font.name
            if hasattr(source_run.font, 'size') and source_run.font.size:
                target_run.font.size = source_run.font.size
    except Exception as e:
        logger.warning(f"Error copying run formatting: {e}")

def copy_paragraph_with_formatting(source_para, target_doc):
    """Copy paragraph with full formatting including hyperlinks"""
    target_para = target_doc.add_paragraph()
    
    try:
        # Handle hyperlinks
        hyperlinks = {}
        for element in source_para._element.xpath('.//w:hyperlink'):
            r_id = element.get(qn('r:id'))
            if r_id and r_id in source_para.part.rels:
                url = source_para.part.rels[r_id].target_ref
                text = ""
                for t_elem in element.xpath('.//w:t'):
                    if t_elem.text:
                        text += t_elem.text
                hyperlinks[id(element)] = {'url': url, 'text': text}
        
        processed_content = set()
        
        for child in source_para._element:
            if child.tag == qn('w:hyperlink'):
                child_id = id(child)
                if child_id in hyperlinks and hyperlinks[child_id]['text'] not in processed_content:
                    link_info = hyperlinks[child_id]
                    add_hyperlink(target_para, link_info['url'], link_info['text'])
                    processed_content.add(link_info['text'])
            elif child.tag == qn('w:r'):
                run_text = ""
                for t_elem in child.xpath('./w:t'):
                    if t_elem.text:
                        run_text += t_elem.text
                
                if run_text and run_text not in processed_content:
                    new_run = target_para.add_run(run_text)
                    
                    # Find source run for formatting
                    source_run = None
                    for run in source_para.runs:
                        if run.text == run_text:
                            source_run = run
                            break
                    
                    if source_run:
                        copy_run_formatting(source_run, new_run)
                    
                    processed_content.add(run_text)
                    
    except Exception as e:
        logger.error(f"Error copying paragraph formatting: {e}")
        target_para.clear()
        target_para.add_run(source_para.text)
    
    return target_para

def process_bibliography(input_file_path: str, output_file_path: Optional[str] = None, validate_only: bool = False) -> Dict[str, any]:
    """Process bibliography document"""
    try:
        doc = Document(input_file_path)
        paragraphs = [(para.text.strip(), para) for para in doc.paragraphs if para.text.strip()]
        
        if validate_only:
            return len(paragraphs)
        
        # Remove duplicates and sort
        unique_paragraphs = {}
        for text, para in paragraphs:
            if text.lower() not in unique_paragraphs:
                unique_paragraphs[text.lower()] = para
        
        sorted_paragraphs = sorted(unique_paragraphs.items(), key=lambda x: x[0])
        
        # Create new document
        new_doc = Document()
        for _, para in sorted_paragraphs:
            copy_paragraph_with_formatting(para, new_doc)
        
        if output_file_path:
            new_doc.save(output_file_path)
        
        return {
            "success": True,
            "original_count": len(paragraphs),
            "unique_count": len(sorted_paragraphs),
            "duplicates_removed": len(paragraphs) - len(sorted_paragraphs)
        }
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }